from docx import Document


def extract_docx(filepath):
    """
    Extract text from a DOCX file, including paragraphs and tables.
    """

    # Open explicitly with a context manager and hand Document() the open
    # file object, instead of Document(filepath). Document(filepath) opens
    # the file (it's a zip archive) internally and doesn't reliably close
    # it right away, which on Windows keeps a lock on it -- causing a
    # PermissionError later when upload.py tries to os.remove() the temp
    # file. Same root cause we already fixed in pdf_reader.py.
    try:
        with open(filepath, "rb") as f:
            document = Document(f)

            parts = []

            for paragraph in document.paragraphs:
                if paragraph.text.strip():
                    parts.append(paragraph.text)

            for table in document.tables:
                for row in table.rows:
                    cells_text = [cell.text.strip() for cell in row.cells]
                    cells_text = [cell for cell in cells_text if cell]
                    if cells_text:
                        parts.append(" | ".join(cells_text))

    except Exception as error:
        raise RuntimeError(f"Could not read DOCX file: {error}") from error

    return "\n".join(parts).strip()