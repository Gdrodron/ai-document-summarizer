"""
routes/download.py

All file-download endpoints:
- Summary only (TXT / PDF)      -> unchanged from your original main.py
- Full report (TXT / PDF / DOCX) -> new
- Extracted/original text (TXT)  -> new
"""

import io
from datetime import datetime

from flask import (
    Response,
    current_app,
    request,
    send_file,
)

from utils.pdf_report import generate_pdf_report

from . import main_bp


# ======================
# DOWNLOAD SUMMARY (TXT) -- unchanged
# ======================

@main_bp.route("/download-summary", methods=["POST"])
def download_summary():
    summary = request.form.get("summary", "")

    return Response(
        summary,
        mimetype="text/plain",
        headers={"Content-Disposition": "attachment; filename=summary.txt"},
    )


# ======================
# DOWNLOAD SUMMARY (PDF) -- unchanged
# ======================

@main_bp.route("/download-pdf", methods=["POST"])
def download_pdf():
    try:
        filename = request.form.get("filename", "document")
        summary = request.form.get("summary", "")

        # Build the PDF directly in memory instead of writing it to disk
        # and deleting it afterwards. On Windows, send_file() keeps the
        # file open while streaming the response, so an after_this_request
        # cleanup that tries to os.remove() it fails with:
        #   PermissionError: The process cannot access the file because
        #   it is being used by another process.
        # reportlab's SimpleDocTemplate happily accepts a file-like object
        # (BytesIO) instead of a path, so there's no file to clean up at all.
        buf = io.BytesIO()
        generate_pdf_report(buf, filename, summary)
        buf.seek(0)

        return send_file(
            buf,
            mimetype="application/pdf",
            as_attachment=True,
            download_name="AI_Report.pdf",
        )

    except Exception as e:
        current_app.logger.exception(e)
        return str(e), 500


# ======================
# helpers for the FULL report (new)
# ======================

def _safe_filename(filename: str, suffix: str, ext: str) -> str:
    base = (filename or "document").rsplit(".", 1)[0]
    base = "".join(c for c in base if c.isalnum() or c in (" ", "_", "-")).strip()
    base = base.replace(" ", "_") or "document"
    return f"{base}_{suffix}.{ext}"


def _build_txt(title: str, sections: list[tuple[str, str]]) -> io.BytesIO:
    lines = [title, "=" * len(title), "",
             f"Generated on {datetime.now().strftime('%B %d, %Y %I:%M %p')}", ""]
    for heading, body in sections:
        lines.append(heading.upper())
        lines.append("-" * len(heading))
        lines.append(body.strip() if body and body.strip() else "None detected.")
        lines.append("")
    return io.BytesIO("\n".join(lines).encode("utf-8"))


def _build_pdf_full(title: str, sections: list[tuple[str, str]]) -> io.BytesIO:
    from xml.sax.saxutils import escape
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        topMargin=0.75 * inch, bottomMargin=0.75 * inch,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("ReportTitle", parent=styles["Title"], fontSize=20, spaceAfter=6)
    meta_style = ParagraphStyle("Meta", parent=styles["Normal"], textColor="#666666", fontSize=9, spaceAfter=16)
    heading_style = ParagraphStyle("SectionHeading", parent=styles["Heading2"], spaceBefore=14, spaceAfter=6)
    body_style = ParagraphStyle("Body", parent=styles["Normal"], fontSize=10.5, leading=15)
    bullet_style = ParagraphStyle("Bullet", parent=body_style, leftIndent=14)

    story = [
        Paragraph(escape(title), title_style),
        Paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y %I:%M %p')}", meta_style),
        HRFlowable(width="100%", color="#dddddd", thickness=1),
    ]

    for heading, body in sections:
        story.append(Paragraph(escape(heading), heading_style))
        if not body or not body.strip():
            story.append(Paragraph("<i>None detected.</i>", body_style))
            continue
        # Escape first so any &, <, > in the AI-generated text doesn't break
        # reportlab's XML parser, THEN add the markup we actually want.
        lines = [escape(ln.strip()) for ln in body.split("\n") if ln.strip()]
        if len(lines) > 1:
            for ln in lines:
                story.append(Paragraph(f"&bull;&nbsp;&nbsp;{ln}", bullet_style))
        else:
            story.append(Paragraph(escape(body).replace("\n", "<br/>"), body_style))
        story.append(Spacer(1, 4))

    doc.build(story)
    buf.seek(0)
    return buf


def _build_docx_full(title: str, sections: list[tuple[str, str]]) -> io.BytesIO:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"Generated on {datetime.now().strftime('%B %d, %Y %I:%M %p')}").italic = True

    for heading, body in sections:
        doc.add_heading(heading, level=2)
        if not body or not body.strip():
            doc.add_paragraph("None detected.")
            continue
        lines = [ln.strip() for ln in body.split("\n") if ln.strip()]
        if len(lines) > 1:
            for ln in lines:
                doc.add_paragraph(ln, style="List Bullet")
        else:
            doc.add_paragraph(body)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ======================
# DOWNLOAD FULL REPORT (new: pdf / txt / docx)
# ======================

@main_bp.route("/download-full-report/<fmt>", methods=["POST"])
def download_full_report(fmt):
    filename = request.form.get("filename", "document")
    summary = request.form.get("summary", "")
    key_points = request.form.get("key_points", "")
    action_items = request.form.get("action_items", "")
    keywords = request.form.get("keywords", "")

    sections = [
        ("Summary", summary),
        ("Key Points", key_points),
        ("Action Items", action_items),
        ("Keywords", keywords),
    ]
    title = f"Analysis Report - {filename}"

    try:
        if fmt == "txt":
            buf = _build_txt(title, sections)
            return send_file(
                buf, mimetype="text/plain", as_attachment=True,
                download_name=_safe_filename(filename, "full_report", "txt"),
            )

        elif fmt == "pdf":
            buf = _build_pdf_full(title, sections)
            return send_file(
                buf, mimetype="application/pdf", as_attachment=True,
                download_name=_safe_filename(filename, "full_report", "pdf"),
            )

        elif fmt == "docx":
            buf = _build_docx_full(title, sections)
            return send_file(
                buf,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True,
                download_name=_safe_filename(filename, "full_report", "docx"),
            )

        else:
            return "Unsupported format.", 400

    except Exception as e:
        current_app.logger.exception(e)
        return str(e), 500


# ======================
# DOWNLOAD EXTRACTED TEXT (new)
# ======================

@main_bp.route("/download-extracted-text", methods=["POST"])
def download_extracted_text():
    filename = request.form.get("filename", "document")
    extracted_text = request.form.get("extracted_text", "")

    return Response(
        extracted_text,
        mimetype="text/plain",
        headers={
            "Content-Disposition":
                f"attachment; filename={_safe_filename(filename, 'extracted', 'txt')}"
        },
    )